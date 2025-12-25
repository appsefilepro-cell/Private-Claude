#!/usr/bin/env python3
"""
PROBATE ADMINISTRATION AUTOMATION SYSTEM
Automates estate administration documents and notifications

Features:
- Letters of Administration generation
- Creditor notification letters
- Bank notification letters
- Insurance policy notifications
- Email and Fax delivery
- Case tracking and management

First Client: Thurman Robinson Estate
"""

import json
import os
import smtplib
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger('ProbateAdmin')


class ProbateAdministrator:
    """
    Complete Probate Administration Automation

    Handles:
    1. Letters of Administration
    2. Creditor notifications
    3. Bank notifications
    4. Insurance notifications
    5. Court filings
    6. Asset inventory
    7. Distribution plans
    """

    def __init__(self):
        self.base_path = Path(__file__).parent
        self.templates_path = self.base_path / "templates"
        self.cases_path = self.base_path / "cases"
        self.output_path = self.base_path / "generated_docs"

        # Create directories
        self.templates_path.mkdir(exist_ok=True)
        self.cases_path.mkdir(exist_ok=True)
        self.output_path.mkdir(exist_ok=True)

        logger.info("=" * 70)
        logger.info("⚖️  PROBATE ADMINISTRATION SYSTEM INITIALIZED")
        logger.info("=" * 70)

    def create_new_case(self, case_info: Dict[str, Any]) -> str:
        """
        Create new probate case

        Required info:
        - decedent_name: Name of deceased
        - date_of_death: Date of death
        - case_number: Court case number
        - administrator_name: Personal representative
        - administrator_address: PR address
        - court_name: Probate court name
        - county: County
        - state: State
        """
        case_id = f"PROBATE_{case_info['decedent_name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"

        # Create case folder structure
        case_folder = self.cases_path / case_id
        case_folder.mkdir(exist_ok=True)

        # Create subfolders
        (case_folder / "letters").mkdir(exist_ok=True)
        (case_folder / "filings").mkdir(exist_ok=True)
        (case_folder / "inventory").mkdir(exist_ok=True)
        (case_folder / "distributions").mkdir(exist_ok=True)
        (case_folder / "correspondence").mkdir(exist_ok=True)

        # Save case info
        case_info['case_id'] = case_id
        case_info['created_date'] = datetime.now().isoformat()

        with open(case_folder / "case_info.json", 'w') as f:
            json.dump(case_info, f, indent=2)

        logger.info(f"✅ Created probate case: {case_id}")
        logger.info(f"   Decedent: {case_info['decedent_name']}")
        logger.info(f"   Case #: {case_info.get('case_number', 'TBD')}")

        return case_id

    def generate_letter_of_administration(self, case_id: str) -> str:
        """Generate Letter of Administration document"""
        case_info = self._load_case_info(case_id)

        doc = Document()

        # Header
        header = doc.add_heading('LETTER OF ADMINISTRATION', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Court info
        court_info = doc.add_paragraph()
        court_info.add_run(f"{case_info['court_name']}\n").bold = True
        court_info.add_run(f"{case_info['county']} County, {case_info['state']}\n")
        court_info.add_run(f"Case No.: {case_info.get('case_number', '_____________')}\n")

        doc.add_paragraph()

        # Estate info
        doc.add_paragraph().add_run(f"In the Matter of the Estate of:").bold = True
        doc.add_paragraph(f"{case_info['decedent_name'].upper()}, Deceased")

        doc.add_paragraph()

        # Body
        body = f"""TO WHOM IT MAY CONCERN:

This is to certify that on {case_info.get('appointment_date', '_______________')}, the undersigned was duly appointed and qualified as Personal Representative (Administrator) of the Estate of {case_info['decedent_name']}, deceased, by the {case_info['court_name']}.

The Personal Representative is hereby authorized to act on behalf of the Estate in all matters, including but not limited to:

1. Collecting, managing, and protecting all assets of the Estate
2. Paying valid debts and claims against the Estate
3. Filing all required tax returns
4. Distributing assets to heirs and beneficiaries
5. Representing the Estate in all legal proceedings

This Letter of Administration shall remain in full force and effect until the Estate is closed or until revoked by the Court.

All persons and institutions are requested to recognize the authority of the Personal Representative and to cooperate fully in the administration of this Estate.

Dated: {datetime.now().strftime('%B %d, %Y')}


_________________________________
{case_info['administrator_name']}
Personal Representative

{case_info.get('administrator_address', '')}
"""

        doc.add_paragraph(body)

        # Court seal section
        doc.add_paragraph()
        seal_para = doc.add_paragraph()
        seal_para.add_run("COURT SEAL:\n\n\n").bold = True
        seal_para.add_run("_________________________________\n")
        seal_para.add_run("Clerk of Court")

        # Save
        output_file = self.output_path / f"{case_id}_letter_of_administration.docx"
        doc.save(output_file)

        logger.info(f"✅ Generated Letter of Administration: {output_file}")

        return str(output_file)

    def generate_creditor_notification(self, case_id: str, creditor_info: Dict[str, Any]) -> str:
        """
        Generate creditor notification letter

        creditor_info should include:
        - creditor_name
        - creditor_address
        - account_number (optional)
        - estimated_debt (optional)
        """
        case_info = self._load_case_info(case_id)

        doc = Document()

        # Letterhead
        header = doc.add_paragraph()
        header.add_run(f"{case_info['administrator_name']}\n").bold = True
        header.add_run("Personal Representative\n")
        header.add_run(f"{case_info.get('administrator_address', '')}\n")
        header.add_run(f"{case_info.get('administrator_phone', '')}\n")
        header.add_run(f"{case_info.get('administrator_email', '')}\n")

        doc.add_paragraph()
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()

        # Recipient
        doc.add_paragraph(f"{creditor_info['creditor_name']}\n{creditor_info['creditor_address']}")
        doc.add_paragraph()

        # Subject line
        subject = doc.add_paragraph()
        subject.add_run(f"RE: Estate of {case_info['decedent_name']}, Deceased\n").bold = True
        subject.add_run(f"     Date of Death: {case_info['date_of_death']}\n")
        if creditor_info.get('account_number'):
            subject.add_run(f"     Account Number: {creditor_info['account_number']}\n")

        doc.add_paragraph()

        # Body
        body = f"""Dear Sir or Madam:

This letter serves as formal notification that {case_info['decedent_name']} passed away on {case_info['date_of_death']}. I have been appointed as Personal Representative of the Estate by the {case_info['court_name']} under Case No. {case_info.get('case_number', 'pending')}.

According to our records, the decedent may have had an account or obligation with your company. If this is correct, please provide the following information within 30 days:

1. Current account balance as of the date of death
2. Complete account statements for the past 12 months
3. Any outstanding debts or obligations
4. Instructions for closing the account or settling the debt
5. Required documentation for estate processing

Please send all correspondence and documentation to:

{case_info['administrator_name']}
Personal Representative
{case_info.get('administrator_address', '')}

All claims against the Estate must be filed in accordance with state law. The statutory deadline for filing claims is typically 3-6 months from the date of first publication of notice to creditors (varies by state).

If you have no record of an account for the decedent, please notify me in writing so I may update the Estate records accordingly.

Thank you for your prompt attention to this matter.

Sincerely,


_________________________________
{case_info['administrator_name']}
Personal Representative
Estate of {case_info['decedent_name']}
"""

        doc.add_paragraph(body)

        # Save
        creditor_safe_name = creditor_info['creditor_name'].replace(' ', '_').replace('/', '_')
        output_file = self.output_path / f"{case_id}_creditor_{creditor_safe_name}.docx"
        doc.save(output_file)

        logger.info(f"✅ Generated creditor letter: {creditor_safe_name}")

        return str(output_file)

    def generate_bank_notification(self, case_id: str, bank_info: Dict[str, Any]) -> str:
        """
        Generate bank notification letter

        bank_info should include:
        - bank_name
        - bank_address
        - account_type (checking, savings, etc.)
        - account_number (last 4 digits)
        """
        case_info = self._load_case_info(case_id)

        doc = Document()

        # Letterhead
        header = doc.add_paragraph()
        header.add_run(f"{case_info['administrator_name']}\n").bold = True
        header.add_run("Personal Representative\n")
        header.add_run(f"{case_info.get('administrator_address', '')}\n")
        header.add_run(f"{case_info.get('administrator_phone', '')}\n")

        doc.add_paragraph()
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()

        # Recipient
        doc.add_paragraph(f"{bank_info['bank_name']}\n{bank_info['bank_address']}")
        doc.add_paragraph()

        # Subject
        subject = doc.add_paragraph()
        subject.add_run(f"RE: Notification of Death - Account Holder\n").bold = True
        subject.add_run(f"     Decedent: {case_info['decedent_name']}\n")
        subject.add_run(f"     Date of Death: {case_info['date_of_death']}\n")
        if bank_info.get('account_number'):
            subject.add_run(f"     Account: {bank_info.get('account_type', 'Account')} ending in {bank_info['account_number']}\n")

        doc.add_paragraph()

        # Body
        body = f"""Dear Sir or Madam:

I am writing to inform you that {case_info['decedent_name']} passed away on {case_info['date_of_death']}. I have been appointed as Personal Representative of the Estate by the {case_info['court_name']}.

According to our records, the decedent maintained {bank_info.get('account_type', 'an account')} with your institution. I am requesting the following actions and information:

1. FREEZE ALL ACCOUNTS: Please place an immediate freeze on all accounts in the decedent's name to prevent unauthorized transactions.

2. ACCOUNT INFORMATION: Please provide:
   - Complete list of all accounts (checking, savings, CDs, safe deposit boxes)
   - Account balances as of date of death
   - Account statements for the past 12 months
   - Information about any joint account holders or beneficiaries
   - Safe deposit box access procedures

3. REQUIRED DOCUMENTATION: Please advise what documentation is required to:
   - Access account information
   - Close accounts
   - Transfer funds to Estate account
   - Access safe deposit box (if applicable)

ENCLOSED DOCUMENTS (please attach):
- Certified copy of Death Certificate
- Letter of Administration (when issued by Court)
- Court Order appointing Personal Representative

Please direct all correspondence to:

{case_info['administrator_name']}
Personal Representative
{case_info.get('administrator_address', '')}
Phone: {case_info.get('administrator_phone', '')}
Email: {case_info.get('administrator_email', '')}

I appreciate your prompt attention to this matter and look forward to working with you to properly administer the Estate.

Sincerely,


_________________________________
{case_info['administrator_name']}
Personal Representative
Estate of {case_info['decedent_name']}


ENCLOSURES:
☐ Certified Death Certificate
☐ Letter of Administration
☐ Court Order
"""

        doc.add_paragraph(body)

        # Save
        bank_safe_name = bank_info['bank_name'].replace(' ', '_').replace('/', '_')
        output_file = self.output_path / f"{case_id}_bank_{bank_safe_name}.docx"
        doc.save(output_file)

        logger.info(f"✅ Generated bank notification: {bank_safe_name}")

        return str(output_file)

    def generate_insurance_notification(self, case_id: str, insurance_info: Dict[str, Any]) -> str:
        """
        Generate insurance company notification

        insurance_info should include:
        - insurance_company
        - company_address
        - policy_type (life, health, auto, etc.)
        - policy_number
        """
        case_info = self._load_case_info(case_id)

        doc = Document()

        # Letterhead
        header = doc.add_paragraph()
        header.add_run(f"{case_info['administrator_name']}\n").bold = True
        header.add_run("Personal Representative\n")
        header.add_run(f"{case_info.get('administrator_address', '')}\n")
        header.add_run(f"{case_info.get('administrator_phone', '')}\n")

        doc.add_paragraph()
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()

        # Recipient
        doc.add_paragraph(f"{insurance_info['insurance_company']}\nClaims Department\n{insurance_info['company_address']}")
        doc.add_paragraph()

        # Subject
        subject = doc.add_paragraph()
        subject.add_run(f"RE: Death Claim Notification\n").bold = True
        subject.add_run(f"     Insured: {case_info['decedent_name']}\n")
        subject.add_run(f"     Date of Death: {case_info['date_of_death']}\n")
        subject.add_run(f"     Policy Type: {insurance_info.get('policy_type', 'Unknown')}\n")
        if insurance_info.get('policy_number'):
            subject.add_run(f"     Policy Number: {insurance_info['policy_number']}\n")

        doc.add_paragraph()

        # Body
        body = f"""Dear Claims Department:

I am writing to notify you of the death of {case_info['decedent_name']}, who passed away on {case_info['date_of_death']}. I have been appointed as Personal Representative of the Estate by the {case_info['court_name']}.

We believe the decedent held a {insurance_info.get('policy_type', 'policy')} with your company. I am requesting the following information and actions:

1. POLICY VERIFICATION:
   - Confirm policy existence and current status
   - Policy type and coverage details
   - Face value / benefit amount
   - Beneficiary designation(s)
   - Premium payment status

2. CLAIMS PROCESS:
   - Required claim forms
   - Required documentation
   - Processing timeline
   - Payment procedures

3. POLICY CANCELLATION (if applicable):
   - For non-life insurance policies (auto, homeowners, etc.)
   - Refund of unused premiums
   - Effective cancellation date

INFORMATION PROVIDED:
- Insured Name: {case_info['decedent_name']}
- Date of Birth: {case_info.get('decedent_dob', 'See Death Certificate')}
- Date of Death: {case_info['date_of_death']}
- Social Security Number: {case_info.get('decedent_ssn', 'Available upon request')}

Please send all claim forms, instructions, and correspondence to:

{case_info['administrator_name']}
Personal Representative
{case_info.get('administrator_address', '')}
Phone: {case_info.get('administrator_phone', '')}
Email: {case_info.get('administrator_email', '')}

ENCLOSED DOCUMENTS:
☐ Certified Death Certificate
☐ Letter of Administration
☐ Claimant Statement Form (if provided)

I appreciate your prompt attention to this claim and look forward to your response within 10 business days.

Sincerely,


_________________________________
{case_info['administrator_name']}
Personal Representative
Estate of {case_info['decedent_name']}
"""

        doc.add_paragraph(body)

        # Save
        insurance_safe_name = insurance_info['insurance_company'].replace(' ', '_')
        output_file = self.output_path / f"{case_id}_insurance_{insurance_safe_name}.docx"
        doc.save(output_file)

        logger.info(f"✅ Generated insurance notification: {insurance_safe_name}")

        return str(output_file)

    def send_via_email(self, document_path: str, recipient_email: str,
                       subject: str = None, body: str = None) -> bool:
        """Send document via email"""
        try:
            # Load SMTP configuration
            smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
            smtp_port = int(os.getenv('SMTP_PORT', '587'))
            smtp_from = os.getenv('SMTP_FROM', 'appsefilepro@gmail.com')
            smtp_password = os.getenv('SMTP_PASSWORD', '')

            if not smtp_password:
                logger.warning("⚠️  SMTP password not configured in .env file")
                logger.info("   Add SMTP_PASSWORD to config/.env to enable email")
                return False

            # Create message
            msg = MIMEMultipart()
            msg['From'] = smtp_from
            msg['To'] = recipient_email
            msg['Subject'] = subject or "Estate Administration Document"

            # Body
            email_body = body or "Please find attached estate administration document."
            msg.attach(MIMEText(email_body, 'plain'))

            # Attach document
            with open(document_path, 'rb') as f:
                attachment = MIMEApplication(f.read(), _subtype='docx')
                attachment.add_header('Content-Disposition', 'attachment',
                                    filename=Path(document_path).name)
                msg.attach(attachment)

            # Send
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_from, smtp_password)
                server.send_message(msg)

            logger.info(f"✅ Email sent to {recipient_email}")
            return True

        except Exception as e:
            logger.error(f"❌ Email send failed: {e}")
            return False

    def send_via_fax(self, document_path: str, fax_number: str) -> bool:
        """
        Send document via fax using email-to-fax gateway

        Popular free/low-cost options:
        1. eFax (efax.com) - Email to: remote_printer.{fax_number}@efaxsend.com
        2. HelloFax (hellosign.com/fax)
        3. FaxZero (faxzero.com) - Free for up to 5 pages
        4. Email-to-Fax via Gmail: {fax_number}@fax.gmail.com (if available)
        """
        try:
            # Clean fax number (remove spaces, dashes, etc.)
            clean_fax = ''.join(filter(str.isdigit, fax_number))

            # Email-to-fax gateways
            fax_gateways = [
                f"{clean_fax}@efaxsend.com",  # eFax
                f"{clean_fax}@fax.gmail.com",  # Gmail Fax (if available)
            ]

            logger.info(f"📠 Sending fax to {fax_number}...")

            # Try each gateway
            for gateway in fax_gateways:
                success = self.send_via_email(
                    document_path=document_path,
                    recipient_email=gateway,
                    subject=f"Fax to {fax_number}",
                    body="Please see attached document for fax transmission."
                )

                if success:
                    logger.info(f"✅ Fax sent via {gateway}")
                    return True

            # If all gateways fail
            logger.warning("⚠️  Fax not sent - configure email-to-fax service")
            logger.info("   Options:")
            logger.info("   1. Sign up for eFax: https://www.efax.com")
            logger.info("   2. Use FaxZero (free): https://faxzero.com")
            logger.info("   3. Use HelloFax: https://www.hellosign.com/fax")

            return False

        except Exception as e:
            logger.error(f"❌ Fax send failed: {e}")
            return False

    def _load_case_info(self, case_id: str) -> Dict[str, Any]:
        """Load case information"""
        case_file = self.cases_path / case_id / "case_info.json"

        if not case_file.exists():
            raise FileNotFoundError(f"Case not found: {case_id}")

        with open(case_file, 'r') as f:
            return json.load(f)


def main():
    """Demo of Probate Administrator"""
    print("""
    ╔═══════════════════════════════════════════════════════════════════╗
    ║          PROBATE ADMINISTRATION AUTOMATION SYSTEM                 ║
    ║              Estate Administration Made Simple                    ║
    ╚═══════════════════════════════════════════════════════════════════╝
    """)

    admin = ProbateAdministrator()

    # Example: Create case for Thurman Robinson
    print("\n📋 Example: Creating Probate Case")
    print("=" * 70)

    thurman_case = {
        'decedent_name': 'Thurman Robinson',
        'date_of_death': '2024-XX-XX',  # Replace with actual date
        'decedent_dob': '19XX-XX-XX',  # Replace with actual DOB
        'decedent_ssn': 'XXX-XX-XXXX',  # Replace with actual SSN (keep confidential)
        'case_number': 'TBD',  # Will be assigned by court
        'administrator_name': 'Your Name Here',  # Replace with your name
        'administrator_address': 'Your Address Here',  # Replace with your address
        'administrator_phone': 'Your Phone Here',
        'administrator_email': 'appsefilepro@gmail.com',
        'court_name': 'Superior Court',  # Replace with actual court name
        'county': 'Your County',  # Replace with actual county
        'state': 'Your State',  # Replace with actual state
        'appointment_date': 'TBD'  # Date appointed by court
    }

    print("\n🎯 INSTRUCTIONS:")
    print("=" * 70)
    print("1. Fill in the case information in the code above")
    print("2. Run: admin.create_new_case(thurman_case)")
    print("3. Generate documents:")
    print("   - admin.generate_letter_of_administration(case_id)")
    print("   - admin.generate_creditor_notification(case_id, creditor_info)")
    print("   - admin.generate_bank_notification(case_id, bank_info)")
    print("   - admin.generate_insurance_notification(case_id, insurance_info)")
    print()
    print("4. Send documents:")
    print("   - admin.send_via_email(document_path, email)")
    print("   - admin.send_via_fax(document_path, fax_number)")
    print("=" * 70)


if __name__ == "__main__":
    main()
